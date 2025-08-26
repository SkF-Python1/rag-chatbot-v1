"""Document processing utilities for RAG system."""

import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import html2text
from loguru import logger

from config import settings


class SimpleTextSplitter:
    """Simple text splitter that mimics RecursiveCharacterTextSplitter."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200, separators: Optional[List[str]] = None):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", " ", ""]
    
    def split_text(self, text: str) -> List[str]:
        """Split text into chunks."""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        for separator in self.separators:
            if separator in text:
                parts = text.split(separator)
                current_chunk = ""
                
                for part in parts:
                    if len(current_chunk) + len(part) + len(separator) <= self.chunk_size:
                        current_chunk += part + separator
                    else:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                            # Add overlap
                            overlap_start = max(0, len(current_chunk) - self.chunk_overlap)
                            current_chunk = current_chunk[overlap_start:] + part + separator
                        else:
                            current_chunk = part + separator
                
                if current_chunk:
                    chunks.append(current_chunk.strip())
                return chunks
        
        # Fallback: split by character count
        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
            chunks.append(text[i:i + self.chunk_size])
        
        return chunks


@dataclass
class DocumentChunk:
    """Represents a chunk of processed document."""
    content: str
    metadata: Dict[str, Any]
    chunk_id: str
    source_file: str
    chunk_index: int


class DocumentProcessor:
    """Handles processing of various document formats."""
    
    def __init__(self):
        """Initialize the document processor."""
        self.text_splitter = SimpleTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
        self.html_converter = html2text.HTML2Text()
        self.html_converter.ignore_links = True
        self.html_converter.ignore_images = True
        
    def process_file(self, file_path: str) -> List[DocumentChunk]:
        """
        Process a single file and return document chunks.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            List of DocumentChunk objects
        """
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path_obj}")
            
        logger.info(f"Processing file: {file_path_obj}")
        
        # Extract text based on file type
        file_extension = file_path_obj.suffix.lower()
        
        if file_extension == '.pdf':
            text = self._extract_pdf_text(file_path_obj)
        elif file_extension in ['.html', '.htm']:
            text = self._extract_html_text(file_path_obj)
        elif file_extension in ['.docx', '.doc']:
            text = self._extract_docx_text(file_path_obj)
        elif file_extension == '.txt':
            text = self._extract_txt_text(file_path_obj)
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            return []
            
        if not text.strip():
            logger.warning(f"No text extracted from: {file_path}")
            return []
            
        # Clean and split text into chunks
        cleaned_text = self._clean_text(text)
        chunks = self.text_splitter.split_text(cleaned_text)
        
        # Create DocumentChunk objects
        document_chunks = []
        for i, chunk in enumerate(chunks):
            if chunk.strip():  # Only include non-empty chunks
                chunk_id = f"{file_path_obj.stem}_{i}"
                metadata = {
                    "source": str(file_path_obj),
                    "file_type": file_extension,
                    "file_size": file_path_obj.stat().st_size,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_name": file_path_obj.name
                }
                
                document_chunks.append(DocumentChunk(
                    content=chunk,
                    metadata=metadata,
                    chunk_id=chunk_id,
                    source_file=str(file_path_obj),
                    chunk_index=i
                ))
                
        logger.info(f"Created {len(document_chunks)} chunks from {file_path_obj}")
        return document_chunks
    
    def process_directory(self, directory_path: str) -> List[DocumentChunk]:
        """
        Process all supported files in a directory.
        
        Args:
            directory_path: Path to the directory to process
            
        Returns:
            List of all DocumentChunk objects from the directory
        """
        directory_path_obj = Path(directory_path)
        
        if not directory_path_obj.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path_obj}")
            
        supported_extensions = {'.pdf', '.html', '.htm', '.docx', '.doc', '.txt'}
        all_chunks = []
        
        for file_path in directory_path_obj.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    chunks = self.process_file(str(file_path))
                    all_chunks.extend(chunks)
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    
        logger.info(f"Processed {len(all_chunks)} total chunks from {directory_path_obj}")
        return all_chunks
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {e}")
            raise
        return text
    
    def _extract_html_text(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            # Use BeautifulSoup for better HTML parsing
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Convert to text using html2text for better formatting
            text = self.html_converter.handle(str(soup))
            return text
            
        except Exception as e:
            logger.error(f"Error extracting HTML text from {file_path}: {e}")
            raise
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(str(file_path))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {e}")
            raise
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text.
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-\'"\/]', '', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Trim whitespace
        text = text.strip()
        
        return text
    
    def get_file_stats(self, file_path: str) -> Dict[str, Any]:
        """
        Get statistics about a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with file statistics
        """
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            return {}
            
        try:
            chunks = self.process_file(file_path)
            return {
                "file_name": file_path_obj.name,
                "file_size": file_path_obj.stat().st_size,
                "file_type": file_path_obj.suffix.lower(),
                "total_chunks": len(chunks),
                "total_characters": sum(len(chunk.content) for chunk in chunks),
                "average_chunk_size": sum(len(chunk.content) for chunk in chunks) / len(chunks) if chunks else 0
            }
        except Exception as e:
            logger.error(f"Error getting file stats for {file_path_obj}: {e}")
            return {}


# Utility function for easy import
def create_document_processor() -> DocumentProcessor:
    """Create and return a DocumentProcessor instance."""
    return DocumentProcessor() 